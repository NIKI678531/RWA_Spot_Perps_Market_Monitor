"""The docx analysis report.

The xlsx answers "what are the numbers". This answers "what should we do about
them", and its first substantive chapter is the alert summary — the system exists to
find products that suddenly started being bought, so that finding leads.

Prose is generated from the same dataset as the workbook, so the two can never
disagree about a figure.
"""

from __future__ import annotations

from datetime import datetime
from decimal import Decimal
from io import BytesIO
from typing import Any, Sequence

from docx import Document
from docx.shared import Pt
from sqlalchemy.orm import Session

from app.core.metrics import MetricScope
from app.models.enums import AlertStatus
from app.services.analytics import concentration
from app.services.report.dataset import (
    ReportDataset,
    amount_of,
    group_sum,
    load,
    sort_by_amount,
)

#: Rows per table. A management report that runs to 200 rows of detail is not read;
#: the xlsx carries the full population.
_TOP_N = 10

_NOT_VERIFIED = "未验证"


def build_docx(session: Session, as_of: datetime | None = None) -> bytes:
    """Render the analysis report for one snapshot."""
    return render_docx(load(session, as_of))


def render_docx(data: ReportDataset) -> bytes:
    document = Document()
    document.add_heading("RWA 代币化资产市场监控报告", level=0)
    intro = document.add_paragraph()
    intro.add_run(f"数据时点：{_ts(data.as_of)}").bold = True
    document.add_paragraph(
        "本报告的五类金额指标——现货市值、现货成交、DEX 流动性、永续成交、永续未平仓——"
        "为互不相容的口径，全文并列展示，任何位置均不做跨口径加总。"
        f"缺失观测标注为「{_NOT_VERIFIED}」，与 0 是两回事。"
    )

    _chapter_alerts(document, data)
    _chapter_scale(document, data)
    _chapter_venues(document, data)
    _chapter_perps(document, data)
    _chapter_quality(document, data)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _chapter_alerts(document: Any, data: ReportDataset) -> None:
    """The alert summary. First chapter by design, not by accident."""
    document.add_heading("一 · 异常告警摘要", level=1)

    if not data.alerts:
        document.add_paragraph(
            "本时点无未处理告警。基线需要 14 个同 session 快照才会开始触发，"
            "系统上线初期无告警属于预期，而非无异常。"
        )
        return

    confirmed = [a for a in data.alerts if a.status is AlertStatus.CONFIRMED]
    document.add_paragraph(
        f"未处理告警 {len(data.alerts)} 条，其中已确认（连续两个快照成立）{len(confirmed)} 条。"
        "全部告警均已通过 5 万美元绝对金额下限——从 500 美元涨到 5,000 美元是 +900%，"
        "但没有商业意义。"
    )

    table = _table(
        document,
        ["严重度", "检测器", "实体", "口径", "结论", "首现", "次数"],
    )
    for alert in data.alerts[:_TOP_N]:
        _row(
            table,
            [
                str(alert.severity.value),
                alert.detector,
                f"{alert.entity_type.value}:{alert.entity_id}",
                str(alert.metric_scope.value),
                alert.headline_zh,
                _ts(alert.first_seen_ts),
                str(alert.occurrence_count),
            ],
        )

    document.add_paragraph(
        "每条告警的原始值、基线中位数、MAD、样本量、market_session 与规则名记录在 "
        "20_Alert_Evidence 表。无法向管理层复述依据的告警等同于噪音。"
    )


def _chapter_scale(document: Any, data: ReportDataset) -> None:
    document.add_heading("二 · 市场规模", level=1)

    union = [c for c in data.categories if c.is_additive]
    overlapping = [c for c in data.categories if not c.is_additive]
    if union:
        row = union[0]
        document.add_paragraph(
            f"去重后合计：资产 {row.asset_count or 0} 个，"
            f"市值 {_usd(row.market_cap)}，24 小时成交 {_usd(row.vol_24h)}。"
        )
    if overlapping:
        document.add_paragraph(
            f"来源方按 {len(overlapping)} 个类别索引这批资产，类别之间按构造重叠"
            "（同一枚代币可同时属于 Tokenized Stock、xStocks 与 Ondo），"
            "相加约得真实值的 2.7 倍。上文只引用去重并集行。"
        )

    demand_spot = group_sum(
        data.scoped_pairs,
        lambda p: p.asset.underlying_id,
        lambda p: p.snapshot.adjusted_vol_24h,
        MetricScope.SPOT_VOLUME,
    )
    if not demand_spot:
        return

    names = {u.underlying_id: u.name for u in data.underlyings}
    document.add_heading("需求最集中的底层", level=2)
    table = _table(document, ["排名", "底层", "名称", "现货成交（质量调整）"])
    for rank, key in enumerate(
        sort_by_amount(list(demand_spot), demand_spot)[:_TOP_N], start=1
    ):
        _row(
            table,
            [
                str(rank),
                key,
                names.get(key, ""),
                _usd(amount_of(demand_spot, key)),
            ],
        )


def _chapter_venues(document: Any, data: ReportDataset) -> None:
    document.add_heading("三 · 场所格局", level=1)

    adjusted = group_sum(
        data.scoped_pairs,
        lambda p: p.snapshot.venue_id,
        lambda p: p.snapshot.adjusted_vol_24h,
        MetricScope.SPOT_VOLUME,
    )
    if not adjusted:
        document.add_paragraph(f"本时点无现货交易对观测（{_NOT_VERIFIED}）。")
        return

    keys = list(adjusted)
    result = concentration.compute([adjusted[k] for k in keys], keys)
    top1 = result.top_n_share(1).value
    top5 = result.top_n_share(5).value
    document.add_paragraph(
        f"共观测到 {len(keys)} 个现货场所，HHI {float(result.hhi):,.0f}"
        f"（0–10000 标度，2500 以上为高集中）。"
        f"第一名占 {_pct(top1)}，前五名合计占 {_pct(top5)}。"
    )

    names = {v.venue_id: v.name for v in data.venues}
    table = _table(document, ["排名", "场所", "质量调整成交", "份额"])
    for rank, share in enumerate(result.shares[:_TOP_N], start=1):
        _row(
            table,
            [
                str(rank),
                names.get(share.entity_id, share.entity_id),
                _usd(share.value),
                _pct(share.share),
            ],
        )


def _chapter_perps(document: Any, data: ReportDataset) -> None:
    document.add_heading("四 · 永续市场", level=1)

    # Scoped, not raw: the collectors enumerate whole exchanges, so an ungated total
    # here would be Hyperliquid's BTC book presented to management as RWA demand.
    contracts = data.scoped_perp_contracts
    if not contracts:
        document.add_paragraph(f"本时点无永续合约观测（{_NOT_VERIFIED}）。")
        return

    volumes = group_sum(
        contracts,
        lambda r: r.snapshot.contract_id,
        lambda r: r.snapshot.vol_24h,
        MetricScope.PERP_VOLUME,
    )
    observed = [v.amount for v in volumes.values() if v.amount is not None]
    total = sum(observed, start=Decimal(0))
    hip3 = {r.snapshot.contract_id for r in contracts if r.perp_dex}

    document.add_paragraph(
        f"共观测到 {len(volumes)} 个映射到真实标的的合约，24 小时成交合计 {_usd(total)}，"
        f"其中 {len(hip3)} 个部署在 HIP-3 独立永续 DEX 上。"
        "聚合站只列 Top 25，看不到无许可部署，因此本表直接向交易所枚举；"
        "交易所自有的加密原生合约不计入本节。"
    )

    by_id = {r.snapshot.contract_id: r for r in contracts}
    table = _table(document, ["排名", "合约", "perp_dex", "成交", "占比"])
    for rank, key in enumerate(
        sort_by_amount(list(volumes), volumes)[:_TOP_N], start=1
    ):
        row = by_id[key]
        amount = amount_of(volumes, key)
        _row(
            table,
            [
                str(rank),
                row.symbol,
                row.perp_dex or "core",
                _usd(amount),
                _pct(amount / total) if amount is not None and total else _NOT_VERIFIED,
            ],
        )

    document.add_paragraph(
        "永续成交与现货成交是两个口径，本报告不给出合计。基线数据集中永续成交约为"
        "现货的 1.7 倍，只看现货会低估真实需求。"
    )


def _chapter_quality(document: Any, data: ReportDataset) -> None:
    document.add_heading("五 · 数据质量", level=1)

    flagged = sum(
        1
        for p in data.pairs
        if p.snapshot.is_quality_anomaly or p.snapshot.is_quality_stale
    )
    unverified = sum(1 for p in data.pairs if p.snapshot.raw_vol_24h is None)
    document.add_paragraph(
        f"现货交易对 {len(data.pairs)} 个，其中报价被来源标记为异常或陈旧的 {flagged} 个，"
        f"成交额未观测到的 {unverified} 个。"
        "原始成交与质量调整成交在 xlsx 中并列，两者可相差三个数量级——"
        "参考案例中某场所原始约 $29.3mn、调整后约 $216，19 个交易对中 17 个被标记。"
    )

    failures = [e for e in data.fetch_log if e.status.value != "ok"]
    if failures:
        document.add_paragraph(
            f"最近的采集记录中有 {len(failures)} 次非成功尝试。"
            f"这些位置在全文标注为「{_NOT_VERIFIED}」，不会被写作 0。"
        )


# --- helpers ---------------------------------------------------------------


def _table(document: Any, headers: Sequence[str]) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    return table


def _row(table: Any, values: Sequence[str]) -> None:
    cells = table.add_row().cells
    for cell, value in zip(cells, values):
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)


def _usd(amount: Decimal | None) -> str:
    """Format money, or say we did not observe it. Never renders a missing value."""
    if amount is None:
        return _NOT_VERIFIED
    value = float(amount)
    for unit, size in (("bn", 1e9), ("mn", 1e6), ("k", 1e3)):
        if abs(value) >= size:
            return f"${value / size:,.2f}{unit}"
    return f"${value:,.2f}"


def _pct(value: Decimal | float | None) -> str:
    if value is None:
        return _NOT_VERIFIED
    return f"{float(value) * 100:.1f}%"


def _ts(value: datetime) -> str:
    return value.strftime("%Y-%m-%d %H:%M %Z").strip()
